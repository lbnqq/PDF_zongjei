#!/usr/bin/env python3
"""
简化版年度总结模板生成器

这个脚本会创建一个基本的年度总结模板.docx文件
包含所有必要的占位符

运行方法：
1. 确保已安装 python-docx: pip install python-docx
2. 运行脚本: python generate_template.py

作者：AI助手
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    def create_template():
        """创建年度总结模板"""
        # 创建新文档
        doc = Document()
        
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
        
        # 标题
        title = doc.add_heading('年度工作总结报告', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 空行
        doc.add_paragraph()
        
        # 基本信息
        doc.add_heading('基本信息', level=2)
        info_table = doc.add_table(rows=2, cols=2)
        info_table.style = 'Table Grid'
        info_table.cell(0, 0).text = '报告人：'
        info_table.cell(0, 1).text = '[[您的姓名]]'
        info_table.cell(1, 0).text = '报告日期：'
        info_table.cell(1, 1).text = '[[报告日期]]'
        
        doc.add_paragraph()
        
        # 各个部分
        sections_data = [
            ('年度总结概述', '[[年度总结概述]]'),
            ('主要成就与贡献', '[[主要成就与贡献]]'),
            ('遇到的挑战及解决方案', '[[遇到的挑战及解决方案]]'),
            ('个人成长与学习', '[[个人成长与学习]]'),
            ('未来展望与计划', '[[未来展望与计划]]')
        ]
        
        for heading, placeholder in sections_data:
            doc.add_heading(heading, level=2)
            doc.add_paragraph(placeholder)
            doc.add_paragraph()  # 空行
        
        # 分隔线和页脚
        separator = doc.add_paragraph('─' * 50)
        separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer = doc.add_paragraph('本报告由AI智能年度总结生成器自动生成')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 保存文件
        filename = '年度总结模板.docx'
        doc.save(filename)
        print(f"✅ 模板文件已创建：{filename}")
        
        return True
    
    if __name__ == "__main__":
        print("🔧 正在创建年度总结模板...")
        if create_template():
            print("🎉 模板创建成功！")
            print("\n📋 包含的占位符：")
            placeholders = [
                "[[您的姓名]]", "[[报告日期]]", "[[年度总结概述]]",
                "[[主要成就与贡献]]", "[[遇到的挑战及解决方案]]",
                "[[个人成长与学习]]", "[[未来展望与计划]]"
            ]
            for p in placeholders:
                print(f"   • {p}")

except ImportError:
    print("❌ 缺少 python-docx 库")
    print("请先安装：pip install python-docx")
    print("然后重新运行此脚本")
except Exception as e:
    print(f"❌ 创建模板失败：{e}")
