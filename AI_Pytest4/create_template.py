#!/usr/bin/env python3
"""
年度总结模板生成器

这个脚本会创建一个符合要求的年度总结模板.docx文件
包含所有必要的占位符，供AI智能年度总结生成器使用

使用方法：
    python create_template.py

作者：AI助手
日期：2025年
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def add_heading_style(doc, style_name, font_size, bold=True, color=None):
    """添加自定义标题样式"""
    try:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(font_size)
        font.bold = bold
        if color:
            font.color.rgb = color
        return style
    except:
        # 如果样式已存在，返回现有样式
        return doc.styles[style_name]

def create_annual_summary_template():
    """创建年度总结模板文档"""
    
    # 创建新的Word文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # 添加自定义样式
    title_style = add_heading_style(doc, 'CustomTitle', 18, True)
    heading_style = add_heading_style(doc, 'CustomHeading', 14, True)
    
    # 文档标题
    title = doc.add_heading('年度工作总结报告', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.name = '微软雅黑'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    
    # 添加空行
    doc.add_paragraph()
    
    # 基本信息表格
    info_heading = doc.add_heading('基本信息', level=2)
    info_heading.runs[0].font.name = '微软雅黑'
    info_heading.runs[0].font.size = Pt(16)
    
    # 创建基本信息表格
    info_table = doc.add_table(rows=2, cols=2)
    info_table.style = 'Table Grid'
    
    # 设置表格内容
    info_table.cell(0, 0).text = '报告人：'
    info_table.cell(0, 1).text = '[[您的姓名]]'
    info_table.cell(1, 0).text = '报告日期：'
    info_table.cell(1, 1).text = '[[报告日期]]'
    
    # 设置表格样式
    for row in info_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.runs[0].font.name = '微软雅黑'
                paragraph.runs[0].font.size = Pt(12)
    
    # 添加空行
    doc.add_paragraph()
    
    # 年度总结概述
    overview_heading = doc.add_heading('年度总结概述', level=2)
    overview_heading.runs[0].font.name = '微软雅黑'
    overview_heading.runs[0].font.size = Pt(16)
    
    overview_para = doc.add_paragraph('[[年度总结概述]]')
    overview_para.runs[0].font.name = '微软雅黑'
    overview_para.runs[0].font.size = Pt(12)
    overview_para.paragraph_format.line_spacing = 1.5
    
    # 添加空行
    doc.add_paragraph()
    
    # 主要成就与贡献
    achievement_heading = doc.add_heading('主要成就与贡献', level=2)
    achievement_heading.runs[0].font.name = '微软雅黑'
    achievement_heading.runs[0].font.size = Pt(16)
    
    achievement_para = doc.add_paragraph('[[主要成就与贡献]]')
    achievement_para.runs[0].font.name = '微软雅黑'
    achievement_para.runs[0].font.size = Pt(12)
    achievement_para.paragraph_format.line_spacing = 1.5
    
    # 添加空行
    doc.add_paragraph()
    
    # 遇到的挑战及解决方案
    challenge_heading = doc.add_heading('遇到的挑战及解决方案', level=2)
    challenge_heading.runs[0].font.name = '微软雅黑'
    challenge_heading.runs[0].font.size = Pt(16)
    
    challenge_para = doc.add_paragraph('[[遇到的挑战及解决方案]]')
    challenge_para.runs[0].font.name = '微软雅黑'
    challenge_para.runs[0].font.size = Pt(12)
    challenge_para.paragraph_format.line_spacing = 1.5
    
    # 添加空行
    doc.add_paragraph()
    
    # 个人成长与学习
    growth_heading = doc.add_heading('个人成长与学习', level=2)
    growth_heading.runs[0].font.name = '微软雅黑'
    growth_heading.runs[0].font.size = Pt(16)
    
    growth_para = doc.add_paragraph('[[个人成长与学习]]')
    growth_para.runs[0].font.name = '微软雅黑'
    growth_para.runs[0].font.size = Pt(12)
    growth_para.paragraph_format.line_spacing = 1.5
    
    # 添加空行
    doc.add_paragraph()
    
    # 未来展望与计划
    future_heading = doc.add_heading('未来展望与计划', level=2)
    future_heading.runs[0].font.name = '微软雅黑'
    future_heading.runs[0].font.size = Pt(16)
    
    future_para = doc.add_paragraph('[[未来展望与计划]]')
    future_para.runs[0].font.name = '微软雅黑'
    future_para.runs[0].font.size = Pt(12)
    future_para.paragraph_format.line_spacing = 1.5
    
    # 添加分隔线
    doc.add_paragraph()
    separator = doc.add_paragraph('─' * 50)
    separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    separator.runs[0].font.name = '微软雅黑'
    separator.runs[0].font.size = Pt(10)
    
    # 添加页脚信息
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('本报告由AI智能年度总结生成器自动生成')
    footer_run.font.name = '微软雅黑'
    footer_run.font.size = Pt(10)
    footer_run.font.italic = True
    
    return doc

def main():
    """主函数"""
    print("🔧 正在创建年度总结模板文件...")
    
    try:
        # 创建模板文档
        doc = create_annual_summary_template()
        
        # 保存文件
        filename = '年度总结模板.docx'
        doc.save(filename)
        
        print(f"✅ 模板文件创建成功：{filename}")
        print("\n📋 模板包含以下占位符：")
        placeholders = [
            "[[您的姓名]]",
            "[[报告日期]]", 
            "[[年度总结概述]]",
            "[[主要成就与贡献]]",
            "[[遇到的挑战及解决方案]]",
            "[[个人成长与学习]]",
            "[[未来展望与计划]]"
        ]
        
        for placeholder in placeholders:
            print(f"   • {placeholder}")
        
        print(f"\n🎯 模板文件已保存到当前目录")
        print("现在您可以启动AI智能年度总结生成器了！")
        
    except Exception as e:
        print(f"❌ 创建模板文件失败：{e}")
        print("请确保已安装python-docx库：pip install python-docx")

if __name__ == "__main__":
    main()
