"""
AI 智能年度总结生成器 - Flask后端应用

这个文件是整个应用的核心，包含以下主要功能：
1. 与科大讯飞星火大模型的WebSocket通信
2. 用户输入处理（文本和文件上传）
3. Word文档模板填充和生成
4. Web API接口提供

作者：AI助手
日期：2025年
"""

# 导入必要的Python标准库
import websocket          # WebSocket客户端，用于与星火大模型通信
import datetime          # 日期时间处理
import hashlib           # 哈希算法，用于API签名
import base64            # Base64编码，用于API认证
import hmac              # HMAC签名算法
import json              # JSON数据处理
from urllib.parse import urlencode, quote_plus  # URL编码工具
import ssl               # SSL安全连接
import io                # 内存中的文件操作
import os                # 操作系统接口，用于环境变量

# 导入第三方库
from dotenv import load_dotenv  # 加载.env环境变量文件
from flask import Flask, request, jsonify, send_file, render_template  # Flask Web框架
from docx import Document       # python-docx库，用于处理Word文档
from docx.shared import Inches  # Word文档尺寸设置（虽然当前未直接使用，但为扩展预留）

# 加载.env文件中的环境变量到系统环境中
# 这样我们就可以通过os.getenv()读取配置信息，而不需要在代码中硬编码敏感信息
load_dotenv()

# ==================== 配置部分 ====================
# 从环境变量获取科大讯飞星火大模型的API凭证
# 这些信息需要在讯飞开放平台(www.xfyun.cn)注册应用后获得
APPID = os.getenv("SPARK_APPID")           # 应用ID
APISecret = os.getenv("SPARK_APISECRET")   # API密钥，用于签名
APIKey = os.getenv("SPARK_APIKEY")         # API密钥，用于身份验证

# 检查必要的API凭证是否已配置
def check_api_credentials():
    """检查API凭证配置"""
    protocol = os.getenv("API_PROTOCOL", "HTTP").upper()

    if protocol == "HTTP":
        # 检查HTTP协议凭证
        http_password = os.getenv("SPARK_HTTP_API_PASSWORD")
        if not http_password:
            print("❌ 错误：缺少HTTP协议API凭证配置")
            print("请确保已设置以下环境变量：")
            print("  - SPARK_HTTP_API_PASSWORD")
            print("可以通过以下方式配置：")
            print("1. 运行配置向导: python setup_config.py")
            print("2. 手动创建.env文件并填入凭证")
            print("3. 访问控制台获取APIpassword: https://console.xfyun.cn/services/bmx1")
            return False
        else:
            print("✅ HTTP协议API凭证配置检查通过")
            return True
    else:
        # 检查WebSocket协议凭证
        if not APPID or not APISecret or not APIKey:
            print("❌ 错误：缺少WebSocket协议API凭证配置")
            print("请确保已设置以下环境变量：")
            print("  - SPARK_APPID")
            print("  - SPARK_APISECRET")
            print("  - SPARK_APIKEY")
            print("可以通过以下方式配置：")
            print("1. 运行配置向导: python setup_config.py")
            print("2. 手动创建.env文件并填入凭证")
            print("3. 参考.env.example文件")
            return False
        else:
            print("✅ WebSocket协议API凭证配置检查通过")
            return True

# 执行凭证检查
credentials_ok = check_api_credentials()

# 星火大模型API的服务器配置（可选，有默认值）
SPARK_HOST = os.getenv("SPARK_HOST", "spark-api.xf-yun.com")      # API主机地址
SPARK_DOMAIN = os.getenv("SPARK_DOMAIN", "generalv3.5")           # 模型版本域名
SPARK_API_PATH = os.getenv("SPARK_API_PATH", "/v3.5/chat")        # API路径
SPARK_URL = f"wss://{SPARK_HOST}{SPARK_API_PATH}"                 # 完整的WebSocket URL

print(f"🔗 API配置: {SPARK_HOST}{SPARK_API_PATH} (域名: {SPARK_DOMAIN})")

# ==================== 工具函数 ====================
def generate_spark_auth_url(host, path, api_key, api_secret):
    """
    生成科大讯飞星火大模型的签名认证URL
    
    这个函数实现了讯飞API要求的签名认证流程：
    1. 获取当前GMT时间
    2. 构造待签名字符串
    3. 使用HMAC-SHA256算法签名
    4. Base64编码
    5. 构造最终的认证URL
    
    参数:
        host: API主机地址
        path: API路径
        api_key: API密钥
        api_secret: API密钥
    
    返回:
        完整的带认证参数的WebSocket URL
    """
    # 获取当前时间，必须是GMT格式（格林威治标准时间）
    now = datetime.datetime.now(datetime.timezone.utc)
    date = now.strftime('%a, %d %b %Y %H:%M:%S GMT')

    # 构造待签名的字符串，格式固定
    signature_origin = f"host: {host}\ndate: {date}\nGET {path} HTTP/1.1"

    # 使用HMAC-SHA256算法对待签名字符串进行签名
    signature_sha256 = hmac.new(
        api_secret.encode('utf-8'),           # 密钥
        signature_origin.encode('utf-8'),     # 待签名数据
        hashlib.sha256                        # 签名算法
    ).digest()
    
    # 将签名结果进行Base64编码
    signature_base64 = base64.b64encode(signature_sha256).decode('utf-8')

    # 构造Authorization头部信息
    authorization_origin = (
        f'api_key="{api_key}", '
        f'algorithm="hmac-sha256", '
        f'headers="host date request-line", '
        f'signature="{signature_base64}"'
    )
    
    # 对Authorization信息进行Base64编码
    authorization_base64 = base64.b64encode(authorization_origin.encode('utf-8')).decode('utf-8')

    # 构造URL参数
    params = {
        "host": host,
        "date": date,
        "authorization": authorization_base64
    }
    
    # 返回完整的认证URL
    return f"{SPARK_URL}?{urlencode(params)}"


# ==================== WebSocket客户端类 ====================
class SparkWebSocketClient:
    """
    科大讯飞星火大模型WebSocket客户端类
    
    这个类封装了与星火大模型的WebSocket通信逻辑，包括：
    - 连接建立和认证
    - 消息发送和接收
    - 响应内容拼接
    - 错误处理
    """
    
    def __init__(self, appid, api_key, api_secret, domain, host, api_path):
        """
        初始化WebSocket客户端
        
        参数:
            appid: 应用ID
            api_key: API密钥
            api_secret: API密钥
            domain: 模型域名
            host: 主机地址
            api_path: API路径
        """
        self.appid = appid
        self.api_key = api_key
        self.api_secret = api_secret
        self.domain = domain
        self.host = host
        self.api_path = api_path
        
        # 用于存储AI响应的变量
        self.result_content = ""      # 拼接AI的完整响应内容
        self.is_completed = False     # 标记响应是否完成
        self.error_message = None     # 存储错误信息
        self.user_input_prompt = ""   # 存储用户输入的提示词

    def on_open(self, ws):
        """
        WebSocket连接建立时的回调函数
        
        当WebSocket连接成功建立后，这个函数会被自动调用
        主要任务是构造请求消息并发送给星火大模型
        """
        print("WebSocket连接已建立，正在发送请求...")
        
        # 构建发送给星火大模型的请求消息（JSON格式）
        message_json = {
            "header": {
                "app_id": self.appid,  # 应用ID
            },
            "parameter": {
                "chat": {
                    "domain": self.domain,      # 模型版本
                    "temperature": 0.5,         # 控制生成内容的随机性，0-1之间，越小越确定
                    "max_tokens": 4096,         # 最大生成Token数量
                    "top_k": 4,                # 从k个最可能的词中选择
                }
            },
            "payload": {
                "message": {
                    "text": [
                        # 系统角色：定义AI助手的身份和任务
                        {
                            "role": "system", 
                            "content": (
                                "你是一个专业的企业年度总结报告智能助手。"
                                "你的任务是根据用户提供的原始工作内容，"
                                "提炼并总结出年度总结报告的关键信息，"
                                "并以结构化的JSON格式输出。"
                                "确保内容真实、简洁、客观。"
                                "如果某个部分信息不足，请留空或简要说明。"
                                "对于列表形式的字段，请输出JSON数组。"
                            )
                        },
                        # 用户角色：包含用户输入和具体的JSON格式要求
                        {
                            "role": "user", 
                            "content": self.user_input_prompt
                        }
                    ]
                }
            }
        }
        
        # 将消息转换为JSON字符串并发送
        try:
            message_str = json.dumps(message_json, ensure_ascii=False)
            print(f"发送请求消息: {message_str[:200]}...")  # 只打印前200个字符
            ws.send(message_str)
            print("请求消息发送成功")
        except Exception as e:
            self.error_message = f"发送请求消息失败: {str(e)}"
            print(self.error_message)
            self.is_completed = True
            ws.close()

    def on_message(self, ws, message):
        """
        接收到WebSocket消息时的回调函数

        星火大模型会分多次发送响应内容，这个函数负责：
        1. 解析每次收到的消息
        2. 检查是否有错误
        3. 拼接响应内容
        4. 判断是否响应完成
        """
        try:
            # 打印原始响应消息用于调试
            print(f"收到WebSocket消息: {message}")

            # 解析收到的JSON消息
            response_data = json.loads(message)

            # 检查响应数据结构
            if 'header' not in response_data:
                self.error_message = "响应数据缺少header字段"
                print(f"错误: {self.error_message}")
                print(f"响应数据结构: {response_data}")
                self.is_completed = True
                ws.close()
                return

            header = response_data['header']

            # 检查API调用是否成功（错误码为0表示成功）
            if header.get('code', -1) != 0:
                error_code = header.get('code', '未知')
                error_msg = header.get('message', '未知错误')
                self.error_message = f"API请求失败，错误码: {error_code}, 错误信息: {error_msg}"
                print(self.error_message)
                self.is_completed = True
                ws.close()
                return

            # 检查payload字段是否存在
            if 'payload' not in response_data:
                self.error_message = "响应数据缺少payload字段"
                print(f"错误: {self.error_message}")
                print(f"响应数据结构: {response_data}")
                self.is_completed = True
                ws.close()
                return

            payload = response_data['payload']

            # 检查choices字段是否存在
            if 'choices' not in payload:
                self.error_message = "payload中缺少choices字段"
                print(f"错误: {self.error_message}")
                print(f"payload结构: {payload}")
                self.is_completed = True
                ws.close()
                return

            # 获取响应内容
            choices = payload['choices']

            # 检查status字段
            if 'status' not in choices:
                self.error_message = "choices中缺少status字段"
                print(f"错误: {self.error_message}")
                print(f"choices结构: {choices}")
                self.is_completed = True
                ws.close()
                return

            status = choices['status']  # 0：开始；1：进行中；2：结束

            # 检查text字段
            if 'text' not in choices or not choices['text']:
                print("警告: choices中缺少text字段或text为空")
                # 不是致命错误，继续处理
            else:
                # 拼接AI生成的内容（星火模型会分多次发送内容）
                if len(choices['text']) > 0 and 'content' in choices['text'][0]:
                    content = choices['text'][0]['content']
                    self.result_content += content
                    print(f"收到内容片段: {content[:100]}...")  # 只打印前100个字符

            # 如果状态为2，表示响应结束
            if status == 2:
                print("AI响应已完成")
                print(f"完整响应内容长度: {len(self.result_content)} 字符")
                self.is_completed = True
                ws.close()  # 关闭WebSocket连接

        except json.JSONDecodeError as e:
            self.error_message = f"JSON解析错误: {str(e)}"
            print(f"JSON解析错误: {e}")
            print(f"原始消息: {message}")
            self.is_completed = True
            ws.close()
        except KeyError as e:
            self.error_message = f"响应数据结构错误，缺少字段: {str(e)}"
            print(f"KeyError: {e}")
            print(f"响应数据: {response_data if 'response_data' in locals() else message}")
            self.is_completed = True
            ws.close()
        except Exception as e:
            self.error_message = f"处理响应消息时出错: {str(e)}"
            print(f"未知错误: {e}")
            print(f"错误类型: {type(e)}")
            print(f"原始消息: {message}")
            self.is_completed = True
            ws.close()

    def on_error(self, ws, error):
        """WebSocket发生错误时的回调函数"""
        self.error_message = f"WebSocket错误: {error}"
        print(self.error_message)
        self.is_completed = True

    def on_close(self, ws, close_status_code, close_msg):
        """WebSocket连接关闭时的回调函数"""
        print(f"WebSocket连接已关闭. 状态码: {close_status_code}, 消息: {close_msg}")
        self.is_completed = True  # 确保即使异常关闭也能标记为完成

    def send_request(self, user_input_text):
        """
        发送请求到星火大模型并等待响应

        这是客户端的主要方法，负责：
        1. 重置状态变量
        2. 生成认证URL
        3. 构造提示词
        4. 建立WebSocket连接
        5. 等待响应完成
        6. 返回结果或抛出异常

        参数:
            user_input_text: 用户输入的原始文本

        返回:
            AI生成的JSON格式响应内容
        """
        # 重置每次请求的状态
        self.result_content = ""
        self.is_completed = False
        self.error_message = None

        # 检查API凭证
        if not self.appid or not self.api_key or not self.api_secret:
            raise Exception("API凭证未配置，请检查环境变量 SPARK_APPID, SPARK_APIKEY, SPARK_APISECRET")

        print(f"📝 用户输入长度: {len(user_input_text)} 字符")

        # 生成带认证信息的WebSocket URL
        try:
            auth_url = generate_spark_auth_url(self.host, self.api_path, self.api_key, self.api_secret)
            print(f"🔗 认证URL生成成功")
        except Exception as e:
            raise Exception(f"生成认证URL失败: {str(e)}")

        # 构造包含JSON格式要求的完整提示词
        self.user_input_prompt = self._create_spark_prompt(user_input_text)
        print(f"📋 提示词构造完成，长度: {len(self.user_input_prompt)} 字符")

        print(f"🌐 正在连接到星火大模型: {self.host}")

        # 创建WebSocket应用实例
        ws = websocket.WebSocketApp(
            auth_url,
            on_open=self.on_open,       # 连接建立时的回调
            on_message=self.on_message, # 收到消息时的回调
            on_error=self.on_error,     # 发生错误时的回调
            on_close=self.on_close      # 连接关闭时的回调
        )

        try:
            # 运行WebSocket连接，直到完成或出错
            # 注意：cert_reqs=ssl.CERT_NONE 仅用于开发和调试
            # 生产环境应该移除这个参数或设置为ssl.CERT_REQUIRED以确保安全
            print("🚀 开始WebSocket连接...")
            ws.run_forever(sslopt={"cert_reqs": ssl.CERT_NONE})
            print("🔚 WebSocket连接已结束")
        except Exception as e:
            self.error_message = f"WebSocket连接异常: {str(e)}"
            print(f"❌ {self.error_message}")

        # 检查是否有错误发生
        if self.error_message:
            raise Exception(self.error_message)

        # 检查是否收到了响应内容
        if not self.result_content.strip():
            raise Exception("未收到AI模型的有效响应内容")

        print(f"✅ 成功收到AI响应，长度: {len(self.result_content)} 字符")
        return self.result_content

    def _create_spark_prompt(self, user_input_text):
        """
        构造发送给星火大模型的提示词

        这个方法创建一个详细的提示词，要求AI：
        1. 分析用户输入的内容
        2. 提取年度总结的关键信息
        3. 按照指定的JSON格式输出结果

        参数:
            user_input_text: 用户输入的原始文本

        返回:
            格式化的提示词字符串
        """
        prompt = f"""
请根据以下用户输入的文本内容，生成一份年度总结报告的关键信息。
如果某个字段没有对应内容，请使用空字符串或空列表。

用户输入内容：
『{user_input_text}』

请严格按照以下JSON格式输出，确保字段名称不变：
{{
  "年度总结概述": "根据上述内容，总结年度工作亮点、整体表现和主要成就，用一句话概括。",
  "主要成就与贡献": [
    "条目1：具体完成了什么，取得了什么成果",
    "条目2：...",
    "..."
  ],
  "遇到的挑战及解决方案": [
    "条目1：遇到了什么困难，如何解决的",
    "条目2：...",
    "..."
  ],
  "个人成长与学习": [
    "条目1：学习了什么新知识/技能，如何应用",
    "条目2：...",
    "..."
  ],
  "未来展望与计划": [
    "条目1：明年的主要工作目标",
    "条目2：...",
    "..."
  ],
  "姓名": "（请根据上下文推断或留空）",
  "报告日期": "（请根据上下文推断或填写当前日期，格式如：YYYY年MM月DD日）"
}}
"""
        return prompt


# ==================== Flask Web应用 ====================
# 创建Flask应用实例，指定模板文件夹位置
app = Flask(__name__, template_folder='templates')

# 初始化星火大模型客户端
# 根据配置选择HTTP或WebSocket协议
def create_spark_client():
    """根据环境变量配置创建合适的星火大模型客户端"""
    protocol = os.getenv("API_PROTOCOL", "HTTP").upper()

    if protocol == "HTTP":
        # 使用HTTP协议（推荐）
        try:
            from spark_http_client import SparkHTTPClient

            api_password = os.getenv("SPARK_HTTP_API_PASSWORD")
            base_url = os.getenv("SPARK_HTTP_BASE_URL", "https://spark-api-open.xf-yun.com/v2")
            model = os.getenv("SPARK_MODEL", "x1")

            if not api_password:
                raise Exception(
                    "HTTP协议需要配置 SPARK_HTTP_API_PASSWORD，"
                    "请在控制台 https://console.xfyun.cn/services/bmx1 获取APIpassword"
                )

            print(f"🔗 使用HTTP协议连接星火大模型X1")
            return SparkHTTPClient(api_password, base_url, model)

        except ImportError:
            print("❌ 无法导入HTTP客户端，回退到WebSocket协议")
            protocol = "WEBSOCKET"

    if protocol == "WEBSOCKET":
        # 使用WebSocket协议（备用）
        print(f"🔗 使用WebSocket协议连接星火大模型")
        return SparkWebSocketClient(
            APPID, APIKey, APISecret,
            SPARK_DOMAIN, SPARK_HOST, SPARK_API_PATH
        )

    else:
        raise Exception(f"不支持的协议类型: {protocol}")

# 创建客户端实例
try:
    spark_client = create_spark_client()
except Exception as e:
    print(f"❌ 创建星火大模型客户端失败: {e}")
    spark_client = None


@app.route('/')
def index():
    """
    首页路由

    当用户访问网站根目录时，返回主页面
    这个页面包含用户输入表单和文件上传功能
    """
    return render_template('index.html')


@app.route('/generate_summary', methods=['POST'])
def generate_summary():
    """
    生成年度总结的API接口

    这是应用的核心API，处理流程如下：
    1. 接收用户输入（文本或文件）
    2. 调用星火大模型分析内容
    3. 解析AI返回的JSON数据
    4. 加载Word模板并填充数据
    5. 生成并返回Word文档

    返回:
        成功：Word文档文件
        失败：JSON格式的错误信息
    """
    # 检查客户端是否初始化成功
    if spark_client is None:
        return jsonify({
            "error": "星火大模型客户端初始化失败，请检查API配置。"
                    "如使用HTTP协议，请确保配置了SPARK_HTTP_API_PASSWORD；"
                    "如使用WebSocket协议，请确保配置了APPID、APIKEY、APISECRET。"
        }), 500

    user_input = ""

    # ==================== 第1步：获取和处理用户输入 ====================
    try:
        # 检查是否有文本输入
        if 'text_input' in request.form and request.form['text_input'].strip():
            user_input = request.form['text_input'].strip()
            print("收到文本输入，长度:", len(user_input))

        # 检查是否有文件上传
        elif 'file' in request.files:
            file = request.files['file']

            # 检查文件是否被选择
            if file.filename == '':
                return jsonify({"error": "未选择文件"}), 400

            print(f"收到上传文件: {file.filename}")

            # 处理.txt文件
            if file.filename.endswith('.txt'):
                try:
                    # 读取文本文件内容，假设编码为UTF-8
                    user_input = file.read().decode('utf-8')
                except UnicodeDecodeError:
                    return jsonify({"error": "文件编码错误，请确保为UTF-8格式"}), 400

            # 处理.docx文件
            elif file.filename.endswith('.docx'):
                try:
                    # 使用python-docx库读取Word文档
                    doc = Document(file)

                    # 提取所有段落的文本
                    for para in doc.paragraphs:
                        user_input += para.text + "\n"

                    # 提取表格中的文本（如果有的话）
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                user_input += cell.text + "\n"

                except Exception as e:
                    return jsonify({"error": f"读取Word文件失败: {str(e)}"}), 400
            else:
                return jsonify({"error": "不支持的文件类型，请上传 .txt 或 .docx 文件"}), 400
        else:
            return jsonify({"error": "请至少输入一些内容或上传一个文件"}), 400

        # 检查输入内容是否为空
        if not user_input.strip():
            return jsonify({"error": "输入内容为空，请提供有效信息"}), 400

        print(f"处理后的用户输入长度: {len(user_input)} 字符")

    except Exception as e:
        print(f"处理用户输入时出错: {e}")
        return jsonify({"error": f"处理输入时出错: {str(e)}"}), 500

    # ==================== 第2步：调用星火大模型分析内容 ====================
    try:
        print("正在调用星火大模型分析内容...")

        # 发送请求到星火大模型并获取响应
        spark_json_str = spark_client.send_request(user_input)

        print("收到星火大模型响应，长度:", len(spark_json_str))

        # 有时星火模型会返回markdown格式的代码块，需要清理
        # 例如：```json\n{...}\n``` 需要提取中间的JSON部分
        if spark_json_str.strip().startswith("```json") and spark_json_str.strip().endswith("```"):
            spark_json_str = spark_json_str.strip()[7:-3].strip()
            print("清理了markdown代码块格式")

        # 解析AI返回的JSON数据
        extracted_content = json.loads(spark_json_str)
        print("成功解析AI返回的JSON数据")

    except json.JSONDecodeError as e:
        print(f"JSON解析错误: {e}")
        print(f"原始响应内容: {spark_json_str}")
        return jsonify({"error": "AI模型返回内容格式错误，请稍后重试或优化输入"}), 500

    except Exception as e:
        print(f"调用星火大模型失败: {e}")
        return jsonify({"error": f"AI模型服务调用失败: {str(e)}"}), 500

    # ==================== 第3步：加载Word模板并填充数据 ====================
    try:
        # 检查模板文件是否存在
        template_path = '年度总结模板.docx'
        if not os.path.exists(template_path):
            return jsonify({
                "error": "年度总结模板文件不存在，请确保 '年度总结模板.docx' 在应用根目录"
            }), 500

        print("正在加载Word模板...")

        # 加载Word模板文档
        doc = Document(template_path)

        def replace_placeholder(paragraph, placeholder, value):
            """
            替换段落中的占位符

            这个内部函数负责：
            1. 检查段落中是否包含指定的占位符
            2. 根据值的类型进行格式化（列表转换为带项目符号的文本）
            3. 替换占位符为实际内容

            参数:
                paragraph: Word文档的段落对象
                placeholder: 要替换的占位符（如：[[年度总结概述]]）
                value: 替换的值（可能是字符串或列表）
            """
            # 处理不同类型的值
            if isinstance(value, list):
                # 将列表转换为带项目符号的多行文本
                formatted_value = "\n".join([f"• {item}" for item in value if item and item.strip()])
                if not formatted_value:  # 如果列表为空
                    formatted_value = "暂无相关内容"
            else:
                # 确保值是字符串，如果为空则提供默认值
                formatted_value = str(value or '暂无相关内容')

            # 在段落中查找并替换占位符
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, formatted_value)
                print(f"已替换占位符: {placeholder}")

            # 注意：直接替换paragraph.text可能会丢失原有的格式
            # 更复杂的实现需要遍历paragraph.runs来保持格式
            # 但为了简化，这里采用直接替换的方式

        print("正在填充模板数据...")

        # 遍历文档中的所有段落，查找并替换占位符
        for paragraph in doc.paragraphs:
            replace_placeholder(paragraph, '[[年度总结概述]]', extracted_content.get('年度总结概述'))
            replace_placeholder(paragraph, '[[主要成就与贡献]]', extracted_content.get('主要成就与贡献'))
            replace_placeholder(paragraph, '[[遇到的挑战及解决方案]]', extracted_content.get('遇到的挑战及解决方案'))
            replace_placeholder(paragraph, '[[个人成长与学习]]', extracted_content.get('个人成长与学习'))
            replace_placeholder(paragraph, '[[未来展望与计划]]', extracted_content.get('未来展望与计划'))
            replace_placeholder(paragraph, '[[您的姓名]]', extracted_content.get('姓名', '未填写'))
            replace_placeholder(paragraph, '[[报告日期]]',
                              extracted_content.get('报告日期', datetime.date.today().strftime('%Y年%m月%d日')))

        # 处理表格中的占位符（如果模板中有表格）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholder(paragraph, '[[您的姓名]]', extracted_content.get('姓名', '未填写'))
                        replace_placeholder(paragraph, '[[报告日期]]',
                                          extracted_content.get('报告日期', datetime.date.today().strftime('%Y年%m月%d日')))

        print("模板数据填充完成")

    except Exception as e:
        print(f"处理Word模板时出错: {e}")
        return jsonify({"error": f"处理Word模板失败: {str(e)}"}), 500

    # ==================== 第4步：生成Word文档并返回给用户 ====================
    try:
        print("正在生成最终的Word文档...")

        # 将文档保存到内存中的字节流
        # 这样可以直接返回给用户，而不需要在服务器上创建临时文件
        byte_io = io.BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)  # 将文件指针移到开头，准备读取

        # 构造下载文件名
        # 格式：姓名-年度总结-日期.docx
        summary_name = extracted_content.get('姓名', '用户')
        summary_date = datetime.date.today().strftime('%Y%m%d')
        download_filename = f"{summary_name}-年度总结-{summary_date}.docx"

        print(f"文档生成完成，文件名: {download_filename}")

        # 返回文件给用户下载
        return send_file(
            byte_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,  # 作为附件下载
            download_name=download_filename  # 指定下载时的文件名
        )

    except Exception as e:
        print(f"生成Word文档时出错: {e}")
        return jsonify({"error": f"文档生成失败: {str(e)}"}), 500


# ==================== 应用启动 ====================
if __name__ == '__main__':
    """
    应用启动入口

    注意：这种启动方式仅适用于开发和测试环境
    生产环境应该使用专业的WSGI服务器，如：
    - Gunicorn: gunicorn -w 4 -b 0.0.0.0:5000 app:app
    - uWSGI: uwsgi --http :5000 --wsgi-file app.py --callable app
    """
    print("=" * 50)
    print("AI 智能年度总结生成器启动中...")
    print("请确保已正确配置科大讯飞星火大模型的API凭证")
    print("访问地址: http://localhost:5000")
    print("=" * 50)

    # 启动Flask开发服务器
    # debug=True: 启用调试模式，代码修改后自动重启
    # host='0.0.0.0': 允许外部访问（不仅限于localhost）
    # port=5000: 指定端口号
    app.run(debug=True, host='0.0.0.0', port=5000)
